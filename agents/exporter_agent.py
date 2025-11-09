"""
Exporter Agent
Generates professional outputs: Excel models, PDF reports, IC memos, slides
"""

from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
import pandas as pd
from loguru import logger

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

from config.settings import get_settings
from engines import DCFResult, CCAResult, LBOResult, MergerResult
from utils.llm_client import LLMClient


class ExporterAgent:
    """
    Exporter Agent - generates professional outputs
    
    Responsibilities:
    - Excel models with multiple tabs
    - PDF/DOCX IC memos
    - Tear sheets (1-2 pages)
    - Slide decks
    - Red flag logs
    """
    
    def __init__(self):
        """Initialize exporter agent"""
        self.settings = get_settings()
        self.llm = LLMClient()
        self.outputs_dir = self.settings.outputs_dir
        self.outputs_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Exporter Agent initialized - outputs: {self.outputs_dir}")
    
    def export_excel_model(
        self,
        symbol: str,
        company_name: str,
        dcf_result: Optional[DCFResult] = None,
        cca_result: Optional[CCAResult] = None,
        lbo_result: Optional[LBOResult] = None,
        financials_df: Optional[pd.DataFrame] = None
    ) -> Path:
        """
        Generate Excel model with multiple tabs
        
        Args:
            symbol: Stock symbol
            company_name: Company name
            dcf_result: DCF results
            cca_result: CCA results
            lbo_result: LBO results
            financials_df: Historical financials
            
        Returns:
            Path to Excel file
        """
        if not OPENPYXL_AVAILABLE:
            logger.error("openpyxl not installed")
            return None
        
        logger.info(f"Generating Excel model for {symbol}")
        
        # Create workbook
        wb = openpyxl.Workbook()
        
        # Remove default sheet
        wb.remove(wb.active)
        
        # 1. Summary Tab
        ws_summary = wb.create_sheet("Summary")
        self._create_summary_tab(ws_summary, symbol, company_name, dcf_result, cca_result, lbo_result)
        
        # 2. DCF Tab
        if dcf_result:
            ws_dcf = wb.create_sheet("DCF")
            self._create_dcf_tab(ws_dcf, dcf_result)
        
        # 3. CCA Tab
        if cca_result:
            ws_cca = wb.create_sheet("CCA")
            self._create_cca_tab(ws_cca, cca_result)
        
        # 4. LBO Tab
        if lbo_result:
            ws_lbo = wb.create_sheet("LBO")
            self._create_lbo_tab(ws_lbo, lbo_result)
        
        # 5. Historical Financials Tab
        if financials_df is not None and not financials_df.empty:
            ws_hist = wb.create_sheet("Hist_Financials")
            self._create_financials_tab(ws_hist, financials_df)
        
        # Save file
        filename = f"{symbol}_Valuation_Model_{datetime.now().strftime('%Y%m%d')}.xlsx"
        filepath = self.outputs_dir / filename
        
        wb.save(filepath)
        logger.info(f"Excel model saved: {filepath}")
        
        return filepath
    
    def _create_summary_tab(self, ws, symbol, company_name, dcf_result, cca_result, lbo_result):
        """Create summary tab"""
        # Header
        ws['A1'] = f"{company_name} ({symbol}) - Valuation Summary"
        ws['A1'].font = Font(size=16, bold=True)
        ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        
        row = 4
        
        # DCF Summary
        if dcf_result:
            ws[f'A{row}'] = "DCF Valuation"
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
            
            ws[f'A{row}'] = "Enterprise Value"
            ws[f'B{row}'] = dcf_result.enterprise_value
            ws[f'B{row}'].number_format = '$#,##0'
            row += 1
            
            ws[f'A{row}'] = "Equity Value"
            ws[f'B{row}'] = dcf_result.equity_value
            ws[f'B{row}'].number_format = '$#,##0'
            row += 1
            
            ws[f'A{row}'] = "Value per Share"
            ws[f'B{row}'] = dcf_result.value_per_share
            ws[f'B{row}'].number_format = '$#,##0.00'
            row += 1
            
            ws[f'A{row}'] = "WACC"
            ws[f'B{row}'] = dcf_result.wacc
            ws[f'B{row}'].number_format = '0.00%'
            row += 2
        
        # CCA Summary
        if cca_result:
            ws[f'A{row}'] = "Comparable Company Analysis"
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
            
            ws[f'A{row}'] = "Value per Share (EV/Revenue)"
            ws[f'B{row}'] = cca_result.value_per_share_revenue
            ws[f'B{row}'].number_format = '$#,##0.00'
            row += 1
            
            ws[f'A{row}'] = "Value per Share (EV/EBITDA)"
            ws[f'B{row}'] = cca_result.value_per_share_ebitda
            ws[f'B{row}'].number_format = '$#,##0.00'
            row += 1
            
            ws[f'A{row}'] = "Peer Count"
            ws[f'B{row}'] = cca_result.peer_count
            row += 2
        
        # LBO Summary
        if lbo_result:
            ws[f'A{row}'] = "LBO Analysis"
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
            
            ws[f'A{row}'] = "Equity IRR"
            ws[f'B{row}'] = lbo_result.equity_irr
            ws[f'B{row}'].number_format = '0.0%'
            row += 1
            
            ws[f'A{row}'] = "MoIC"
            ws[f'B{row}'] = lbo_result.equity_moic
            ws[f'B{row}'].number_format = '0.00x'
            row += 1
        
        # Auto-adjust column width
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20
    
    def _create_dcf_tab(self, ws, dcf_result):
        """Create DCF detailed tab"""
        ws['A1'] = "DCF Valuation Details"
        ws['A1'].font = Font(size=14, bold=True)
        
        row = 3
        ws[f'A{row}'] = "WACC Calculation"
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        
        ws[f'A{row}'] = "WACC"
        ws[f'B{row}'] = dcf_result.wacc
        ws[f'B{row}'].number_format = '0.00%'
        row += 1
        
        ws[f'A{row}'] = "Cost of Equity"
        ws[f'B{row}'] = dcf_result.cost_of_equity
        ws[f'B{row}'].number_format = '0.00%'
        row += 1
        
        ws[f'A{row}'] = "Levered Beta"
        ws[f'B{row}'] = dcf_result.levered_beta
        ws[f'B{row}'].number_format = '0.00'
        row += 2
        
        ws[f'A{row}'] = "Valuation Bridge"
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        
        ws[f'A{row}'] = "PV of Forecast Period"
        ws[f'B{row}'] = dcf_result.pv_forecast_period
        ws[f'B{row}'].number_format = '$#,##0'
        row += 1
        
        ws[f'A{row}'] = "PV of Terminal Value"
        ws[f'B{row}'] = dcf_result.pv_terminal_value
        ws[f'B{row}'].number_format = '$#,##0'
        row += 1
        
        ws[f'A{row}'] = "Enterprise Value"
        ws[f'B{row}'] = dcf_result.enterprise_value
        ws[f'B{row}'].number_format = '$#,##0'
        ws[f'B{row}'].font = Font(bold=True)
        
        # Sensitivity table if available
        if dcf_result.sensitivities is not None:
            row += 3
            ws[f'A{row}'] = "Sensitivity Analysis"
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
            
            # Add sensitivity DataFrame
            for r_idx, row_data in enumerate(dataframe_to_rows(dcf_result.sensitivities, index=True, header=True)):
                for c_idx, value in enumerate(row_data):
                    ws.cell(row=row+r_idx, column=1+c_idx, value=value)
        
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20
    
    def _create_cca_tab(self, ws, cca_result):
        """Create CCA detailed tab"""
        ws['A1'] = "Comparable Company Analysis"
        ws['A1'].font = Font(size=14, bold=True)
        
        row = 3
        
        # Summary statistics
        if cca_result.multiples_summary is not None:
            ws[f'A{row}'] = "Multiples Summary"
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
            
            for r_idx, row_data in enumerate(dataframe_to_rows(cca_result.multiples_summary, index=False, header=True)):
                for c_idx, value in enumerate(row_data):
                    cell = ws.cell(row=row+r_idx, column=1+c_idx, value=value)
                    if r_idx == 0:  # Header
                        cell.font = Font(bold=True)
        
        ws.column_dimensions['A'].width = 20
        for col in ['B', 'C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 15
    
    def _create_lbo_tab(self, ws, lbo_result):
        """Create LBO detailed tab"""
        ws['A1'] = "LBO Analysis"
        ws['A1'].font = Font(size=14, bold=True)
        
        row = 3
        ws[f'A{row}'] = "Returns"
        ws[f'A{row}'].font = Font(bold=True)
        row += 1
        
        ws[f'A{row}'] = "Equity IRR"
        ws[f'B{row}'] = lbo_result.equity_irr
        ws[f'B{row}'].number_format = '0.0%'
        row += 1
        
        ws[f'A{row}'] = "MoIC"
        ws[f'B{row}'] = lbo_result.equity_moic
        ws[f'B{row}'].number_format = '0.00x'
        row += 2
        
        # Sources & Uses
        if lbo_result.sources_and_uses is not None:
            ws[f'A{row}'] = "Sources & Uses"
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
            
            for r_idx, row_data in enumerate(dataframe_to_rows(lbo_result.sources_and_uses, index=True, header=True)):
                for c_idx, value in enumerate(row_data):
                    ws.cell(row=row+r_idx, column=1+c_idx, value=value)
        
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20
    
    def _create_financials_tab(self, ws, financials_df):
        """Create historical financials tab"""
        ws['A1'] = "Historical Financials"
        ws['A1'].font = Font(size=14, bold=True)
        
        # Add DataFrame
        for r_idx, row in enumerate(dataframe_to_rows(financials_df, index=False, header=True), start=3):
            for c_idx, value in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)
                if r_idx == 3:  # Header
                    cell.font = Font(bold=True)
        
        # Auto-adjust columns
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 15
    
    def generate_ic_memo(
        self,
        symbol: str,
        company_name: str,
        valuation_summary: Dict[str, Any],
        key_findings: List[str],
        risks: List[str],
        recommendation: str
    ) -> Path:
        """
        Generate Investment Committee Memo (DOCX)
        
        Args:
            symbol: Stock symbol
            company_name: Company name
            valuation_summary: Valuation results
            key_findings: Key findings list
            risks: Risk factors list
            recommendation: Investment recommendation
            
        Returns:
            Path to DOCX file
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            return None
        
        logger.info(f"Generating IC Memo for {symbol}")
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Investment Committee Memorandum", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company info
        doc.add_heading(f"{company_name} ({symbol})", 1)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        
        # Executive Summary
        doc.add_heading("Executive Summary", 2)
        exec_summary = self.llm.write_ic_memo_section(
            section="Executive Summary",
            data=valuation_summary,
            citations=None
        )
        doc.add_paragraph(exec_summary)
        doc.add_paragraph("")
        
        # Valuation Summary
        doc.add_heading("Valuation Summary", 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Method'
        hdr_cells[1].text = 'Value'
        
        for method, value in valuation_summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = method
            row_cells[1].text = f"${value:,.2f}" if isinstance(value, (int, float)) else str(value)
        
        doc.add_paragraph("")
        
        # Key Findings
        doc.add_heading("Key Findings", 2)
        for finding in key_findings:
            doc.add_paragraph(finding, style='List Bullet')
        doc.add_paragraph("")
        
        # Risk Factors
        doc.add_heading("Risk Factors", 2)
        for risk in risks:
            doc.add_paragraph(risk, style='List Bullet')
        doc.add_paragraph("")
        
        # Recommendation
        doc.add_heading("Recommendation", 2)
        doc.add_paragraph(recommendation)
        
        # Save file
        filename = f"{symbol}_IC_Memo_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = self.outputs_dir / filename
        
        doc.save(filepath)
        logger.info(f"IC Memo saved: {filepath}")
        
        return filepath
    
    def generate_tear_sheet(
        self,
        symbol: str,
        company_name: str,
        valuation_range: tuple,
        current_price: float,
        upside_downside: float,
        key_metrics: Dict[str, Any]
    ) -> Path:
        """
        Generate 1-page tear sheet (DOCX)
        
        Args:
            symbol: Stock symbol
            company_name: Company name
            valuation_range: (low, high) tuple
            current_price: Current stock price
            upside_downside: Upside/downside %
            key_metrics: Key financial metrics
            
        Returns:
            Path to DOCX file
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed")
            return None
        
        logger.info(f"Generating tear sheet for {symbol}")
        
        doc = Document()
        
        # Set narrow margins for 1-page format
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Header
        header = doc.add_heading(f"{company_name} ({symbol})", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Valuation box
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = 'Valuation Range'
        cells[1].text = f"${valuation_range[0]:.2f} - ${valuation_range[1]:.2f}"
        
        cells = table.rows[1].cells
        cells[0].text = 'Current Price'
        cells[1].text = f"${current_price:.2f}"
        
        cells = table.rows[2].cells
        cells[0].text = 'Upside/Downside'
        cells[1].text = f"{upside_downside:+.1f}%"
        
        cells = table.rows[3].cells
        cells[0].text = 'Date'
        cells[1].text = datetime.now().strftime('%Y-%m-%d')
        
        doc.add_paragraph("")
        
        # Key Metrics
        doc.add_heading("Key Metrics", 2)
        metrics_table = doc.add_table(rows=len(key_metrics)+1, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        
        hdr_cells = metrics_table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        for idx, (metric, value) in enumerate(key_metrics.items(), start=1):
            row_cells = metrics_table.rows[idx].cells
            row_cells[0].text = metric
            row_cells[1].text = str(value)
        
        # Save file
        filename = f"{symbol}_Tear_Sheet_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = self.outputs_dir / filename
        
        doc.save(filepath)
        logger.info(f"Tear sheet saved: {filepath}")
        
        return filepath


# Example usage
if __name__ == "__main__":
    from engines import DCFResult, CCAResult
    
    # Initialize agent
    agent = ExporterAgent()
    
    print("\n" + "="*70)
    print("EXPORTER AGENT - GENERATING OUTPUTS")
    print("="*70)
    
    symbol = "AAPL"
    company_name = "Apple Inc."
    
    # Mock results (in real use, these come from modeling agent)
    dcf_result = DCFResult(
        enterprise_value=2_900_000_000_000,
        equity_value=2_850_000_000_000,
        value_per_share=183.87,
        shares_outstanding=15_500_000_000,
        pv_forecast_period=1_200_000_000_000,
        terminal_value=2_500_000_000_000,
        pv_terminal_value=1_700_000_000_000,
        wacc=0.078,
        cost_of_equity=0.095,
        levered_beta=1.15,
        cost_of_debt_after_tax=0.030,
        weight_equity=0.95,
        weight_debt=0.05
    )
    
    # 1. Generate Excel Model
    print("\n[1/3] Generating Excel model...")
    excel_path = agent.export_excel_model(
        symbol=symbol,
        company_name=company_name,
        dcf_result=dcf_result
    )
    print(f"✓ Excel saved: {excel_path}")
    
    # 2. Generate IC Memo
    print("\n[2/3] Generating IC Memo...")
    ic_memo_path = agent.generate_ic_memo(
        symbol=symbol,
        company_name=company_name,
        valuation_summary={
            "DCF": 183.87,
            "CCA (EV/EBITDA)": 178.25,
            "CCA (P/E)": 195.50
        },
        key_findings=[
            "Strong cash generation with $165B cash on hand",
            "Services segment growing faster than products",
            "Premium valuation justified by ecosystem moat"
        ],
        risks=[
            "iPhone revenue concentration risk",
            "Regulatory scrutiny on App Store practices",
            "China geopolitical tensions"
        ],
        recommendation="BUY - Target price $185, representing 10% upside from current levels."
    )
    print(f"✓ IC Memo saved: {ic_memo_path}")
    
    # 3. Generate Tear Sheet
    print("\n[3/3] Generating tear sheet...")
    tear_sheet_path = agent.generate_tear_sheet(
        symbol=symbol,
        company_name=company_name,
        valuation_range=(175, 195),
        current_price=178.50,
        upside_downside=3.6,
        key_metrics={
            "Revenue (TTM)": "$383.3B",
            "EBITDA Margin": "32.1%",
            "P/E Ratio": "28.5x",
            "EV/EBITDA": "22.8x",
            "FCF Yield": "3.8%"
        }
    )
    print(f"✓ Tear sheet saved: {tear_sheet_path}")
    
    print("\n" + "="*70)
    print("EXPORTS COMPLETE")
    print("="*70)
